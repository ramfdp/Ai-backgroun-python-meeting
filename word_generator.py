from docx import Document
from docx.shared import Pt
import os

def export_to_word(content, filename="output.docx"):
    # ponytail: they asked for the library, so we add the abstraction.
    doc = Document()
    
    # Basic styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Assuming content might just be a string for now, but we can structure it
    doc.add_heading('Meeting Transcript', level=1)
    doc.add_paragraph(content)
    
    doc.save(filename)
    return os.path.abspath(filename)

if __name__ == "__main__":
    out_path = export_to_word("This is a test document generated with python-docx.\nIt looks neater.")
    print(f"Created {out_path}")
